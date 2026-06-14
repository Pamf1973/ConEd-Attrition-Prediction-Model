#!/usr/bin/env python3
"""
ConEd Steam Attrition Dashboard — Technical Addendum Generator
Covers changes from 2026-06-13 14:08 (manual baseline) through 2026-06-14
Produces: ConEd_Dashboard_Technical_Addendum_Jun2026.docx
"""

import json, math, io, os, datetime
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.gridspec import GridSpec
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE           = os.path.dirname(__file__)
BUILDINGS_FILE = os.path.join(BASE, "public/buildings.json")
ENRICHMENT_FILE= os.path.join(BASE, "public/buildingEnrichment.json")
OUT_FILE       = os.path.join(BASE, "ConEd_Dashboard_Technical_Addendum_Jun2026.docx")

CONED_RED   = RGBColor(0xC8, 0x1D, 0x25)
CONED_NAVY  = RGBColor(0x00, 0x24, 0x69)
SLATE       = RGBColor(0x47, 0x56, 0x69)
ORANGE      = RGBColor(0xE8, 0x77, 0x22)

# ── Load data ──────────────────────────────────────────────────────────────────
with open(BUILDINGS_FILE) as f:
    buildings = json.load(f)
with open(ENRICHMENT_FILE) as f:
    enrichment = json.load(f)

# Merge enrichment onto buildings
enriched = []
for b in buildings:
    key = b.get("address", "").upper()
    e   = enrichment.get(key, {})
    enriched.append({**b, **e, "risk": e.get("ml_risk", b.get("risk"))})

# ── Helpers ────────────────────────────────────────────────────────────────────
def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    buf.seek(0)
    plt.close(fig)
    return buf

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    if level == 1:
        run.font.color.rgb = CONED_NAVY
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = CONED_RED
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = SLATE
        run.font.size = Pt(11)
    return p

def add_body(doc, text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic= italic
    if color:
        run.font.color.rgb = color
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F1F5F9")
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "002469")
        tcPr.append(shading)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri+1]
        fill = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), fill)
            tcPr.append(shading)
    return tbl

def page_break(doc):
    doc.add_page_break()

# ── Charts ─────────────────────────────────────────────────────────────────────

def chart_risk_distribution():
    """Risk distribution — current vs. manual baseline (Uncertain tier removed)."""
    risks = [b["risk"] for b in enriched if b.get("risk") is not None and isinstance(b["risk"], float)]
    bins  = np.arange(0, 1.05, 0.1)
    fig, ax = plt.subplots(figsize=(7, 3.5), facecolor="#030D1A")
    ax.set_facecolor("#001748")
    counts, _ = np.histogram(risks, bins=bins)
    bar_width = (bins[1] - bins[0]) * 0.95
    centers = (bins[:-1] + bins[1:]) / 2
    for i, (c, cnt) in enumerate(zip(centers, counts)):
        if bins[i] >= 0.7:   col = "#ef4444"
        elif bins[i] >= 0.4: col = "#f97316"
        else:                col = "#22c55e"
        ax.bar(c, cnt, width=bar_width, color=col, edgecolor="#0F3B7E", linewidth=0.5)
    ax.axvline(0.4, color="#f97316", linewidth=1.5, linestyle="--", alpha=0.7, label="Medium threshold (0.4)")
    ax.axvline(0.7, color="#ef4444", linewidth=1.5, linestyle="--", alpha=0.7, label="High threshold (0.7)")
    ax.set_xlabel("Attrition Risk Score", color="#94a3b8", fontsize=9)
    ax.set_ylabel("Buildings", color="#94a3b8", fontsize=9)
    ax.set_title("Risk Distribution — 1,210 Buildings (100% ML Coverage)", color="#e2e8f0", fontsize=10, pad=8)
    ax.tick_params(colors="#64748b", labelsize=8)
    for spine in ax.spines.values():
        spine.set_edgecolor("#0F3B7E")
    ax.legend(fontsize=8, facecolor="#001748", edgecolor="#0F3B7E", labelcolor="#94a3b8")
    high   = sum(1 for r in risks if r > 0.7)
    medium = sum(1 for r in risks if 0.4 < r <= 0.7)
    low    = sum(1 for r in risks if r <= 0.4)
    ax.text(0.98, 0.95, f"High: {high}  Medium: {medium}  Low: {low}",
            transform=ax.transAxes, ha="right", va="top",
            color="#e2e8f0", fontsize=8, fontweight="bold",
            bbox=dict(facecolor="#002469", edgecolor="#0F3B7E", boxstyle="round,pad=0.3"))
    fig.tight_layout()
    return fig_to_stream(fig)

def chart_ll97_exposure():
    """LL97 penalty exposure: 2024 vs 2030 by tier."""
    high_24 = sum(b.get("ll97_penalty_2024", 0) or 0 for b in enriched if (b.get("risk") or 0) > 0.7)
    med_24  = sum(b.get("ll97_penalty_2024", 0) or 0 for b in enriched if 0.4 < (b.get("risk") or 0) <= 0.7)
    low_24  = sum(b.get("ll97_penalty_2024", 0) or 0 for b in enriched if (b.get("risk") or 0) <= 0.4)
    high_30 = sum(b.get("ll97_penalty_2030", 0) or 0 for b in enriched if (b.get("risk") or 0) > 0.7)
    med_30  = sum(b.get("ll97_penalty_2030", 0) or 0 for b in enriched if 0.4 < (b.get("risk") or 0) <= 0.7)
    low_30  = sum(b.get("ll97_penalty_2030", 0) or 0 for b in enriched if (b.get("risk") or 0) <= 0.4)

    labels  = ["High Risk", "Medium Risk", "Low Risk"]
    vals_24 = [high_24/1e6, med_24/1e6, low_24/1e6]
    vals_30 = [high_30/1e6, med_30/1e6, low_30/1e6]
    x = np.arange(len(labels))
    w = 0.35

    fig, ax = plt.subplots(figsize=(7, 3.5), facecolor="#030D1A")
    ax.set_facecolor("#001748")
    b1 = ax.bar(x - w/2, vals_24, w, label="2024–2029", color="#3b82f6", edgecolor="#0F3B7E")
    b2 = ax.bar(x + w/2, vals_30, w, label="2030–2034", color="#E87722", edgecolor="#0F3B7E")
    ax.set_ylabel("Annual LL97 Penalty ($M)", color="#94a3b8", fontsize=9)
    ax.set_title("LL97 Penalty Exposure by Risk Tier", color="#e2e8f0", fontsize=10, pad=8)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, color="#94a3b8", fontsize=9)
    ax.tick_params(colors="#64748b", labelsize=8)
    for spine in ax.spines.values(): spine.set_edgecolor("#0F3B7E")
    ax.legend(fontsize=8, facecolor="#001748", edgecolor="#0F3B7E", labelcolor="#94a3b8")
    for bar in list(b1) + list(b2):
        h = bar.get_height()
        if h > 0.1:
            ax.text(bar.get_x() + bar.get_width()/2, h + 0.3, f"${h:.1f}M",
                    ha="center", va="bottom", fontsize=7, color="#e2e8f0")
    fig.tight_layout()
    return fig_to_stream(fig)

def chart_new_endpoints():
    """Architecture diagram — new server endpoints."""
    fig, ax = plt.subplots(figsize=(8, 4), facecolor="#030D1A")
    ax.set_facecolor("#030D1A")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")

    def box(x, y, w, h, label, sublabel="", color="#002469", tc="#e2e8f0", border="#0F3B7E"):
        rect = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor=border, linewidth=1.5, zorder=2)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.15 if sublabel else 0), label,
                ha="center", va="center", color=tc, fontsize=8, fontweight="bold", zorder=3)
        if sublabel:
            ax.text(x + w/2, y + h/2 - 0.25, sublabel,
                    ha="center", va="center", color="#64748b", fontsize=7, zorder=3)

    # Client
    box(0.2, 4.2, 2.0, 1.2, "React Client", "App.jsx + hooks", color="#001748")
    # Server
    box(3.8, 2.0, 4.0, 3.5, "Express API  (api/server.js)", color="#001748")
    # Endpoints
    box(4.2, 4.5, 1.5, 0.7, "GET /api/buildings", "", color="#0041A8", tc="#ffffff")
    box(6.0, 4.5, 1.5, 0.7, "GET /api/watchlist/load", "", color="#0041A8", tc="#ffffff")
    box(4.2, 3.5, 1.5, 0.7, "POST /api/watchlist/save", "", color="#0041A8", tc="#ffffff")
    box(6.0, 3.5, 1.5, 0.7, "GET /api/data/*", "", color="#002469", tc="#94a3b8")
    box(4.2, 2.5, 1.5, 0.7, "POST /api/query", "", color="#002469", tc="#94a3b8")
    box(6.0, 2.5, 1.5, 0.7, "POST /api/auth/*", "", color="#002469", tc="#94a3b8")
    # In-memory store
    box(8.2, 3.8, 1.5, 1.2, "watchlistStore", "Map<token,addr[]>", color="#1e293b", tc="#94a3b8", border="#334155")
    # Arrows
    ax.annotate("", xy=(3.8, 4.8), xytext=(2.2, 4.8),
                arrowprops=dict(arrowstyle="->", color="#0F3B7E", lw=1.5))
    ax.annotate("", xy=(8.2, 4.6), xytext=(7.8, 4.85),
                arrowprops=dict(arrowstyle="->", color="#0F3B7E", lw=1.5))

    ax.text(5.0, 5.7, "New endpoints (blue) — added this sprint", ha="center",
            color="#E87722", fontsize=8, fontweight="bold")
    fig.tight_layout()
    return fig_to_stream(fig)

def chart_pagination_flow():
    """Pagination state machine diagram."""
    fig, ax = plt.subplots(figsize=(8, 2.8), facecolor="#030D1A")
    ax.set_facecolor("#030D1A")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3)
    ax.axis("off")

    nodes = [
        (0.5, 1.5, "All 1,210\nBuildings"),
        (2.5, 1.5, "Apply\nFilters"),
        (4.5, 1.5, "Multi-col\nSort"),
        (6.5, 1.5, "Paginate\n50/page"),
        (8.8, 1.5, "Render\nTable"),
    ]
    for x, y, label in nodes:
        circ = plt.Circle((x, y), 0.6, color="#0041A8", zorder=2)
        ax.add_patch(circ)
        ax.text(x, y, label, ha="center", va="center", color="white",
                fontsize=7, fontweight="bold", zorder=3)

    for i in range(len(nodes)-1):
        x1 = nodes[i][0] + 0.6
        x2 = nodes[i+1][0] - 0.6
        y  = nodes[i][1]
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", color="#E87722", lw=2))

    labels = ["cross-field\nsearch + 7 filters", "sortStack[]\nmulti-key", "page×50\nslice"]
    xpos   = [1.5, 3.5, 5.5]
    for x, lbl in zip(xpos, labels):
        ax.text(x, 0.5, lbl, ha="center", color="#64748b", fontsize=7)

    ax.text(5, 2.7, "RiskTable Rendering Pipeline (new)", ha="center",
            color="#E87722", fontsize=9, fontweight="bold")
    return fig_to_stream(fig)

def chart_watchlist_sync():
    """Watchlist persistence flow."""
    fig, ax = plt.subplots(figsize=(8, 3), facecolor="#030D1A")
    ax.set_facecolor("#030D1A")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis("off")

    def rbox(x, y, w, h, label, color="#001748", tc="#e2e8f0"):
        r = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor="#0F3B7E", lw=1.5, zorder=2)
        ax.add_patch(r)
        ax.text(x+w/2, y+h/2, label, ha="center", va="center",
                color=tc, fontsize=8, fontweight="bold", zorder=3)

    rbox(0.3, 2.2, 2.2, 1.0, "useWatchlist(token)\nWatchlist.jsx")
    rbox(0.3, 0.5, 2.2, 0.8, "localStorage\nconed_watchlist")
    rbox(3.8, 2.2, 2.6, 1.0, "POST /api/watchlist/save\nGET /api/watchlist/load")
    rbox(7.0, 2.2, 2.5, 1.0, "watchlistStore\nMap<token, addr[]>")
    rbox(3.8, 0.5, 2.6, 0.8, "401 if no token\n≤500 sessions, ≤500 chars/addr", color="#1e293b", tc="#94a3b8")

    ax.annotate("", xy=(3.8, 2.7), xytext=(2.5, 2.7),
                arrowprops=dict(arrowstyle="->", color="#E87722", lw=1.5))
    ax.text(3.15, 2.9, "fetch + Bearer token", ha="center", color="#64748b", fontsize=7)
    ax.annotate("", xy=(7.0, 2.7), xytext=(6.4, 2.7),
                arrowprops=dict(arrowstyle="->", color="#E87722", lw=1.5))
    ax.text(6.7, 2.9, "read/write", ha="center", color="#64748b", fontsize=7)
    ax.annotate("", xy=(1.4, 2.2), xytext=(1.4, 1.3),
                arrowprops=dict(arrowstyle="<->", color="#64748b", lw=1.2))
    ax.text(1.7, 1.65, "fallback", fontsize=7, color="#64748b")

    ax.text(5, 3.7, "Watchlist Persistence Architecture (new)", ha="center",
            color="#E87722", fontsize=9, fontweight="bold")
    return fig_to_stream(fig)

def chart_user_interaction_flow():
    """Cross-component interaction flow diagram."""
    fig, ax = plt.subplots(figsize=(9, 4.5), facecolor="#030D1A")
    ax.set_facecolor("#030D1A")
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 6.5)
    ax.axis("off")

    def rbox(x, y, w, h, label, color="#001748", tc="#e2e8f0"):
        r = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor="#0F3B7E", lw=1.5, zorder=2)
        ax.add_patch(r)
        ax.text(x+w/2, y+h/2, label, ha="center", va="center",
                color=tc, fontsize=8, fontweight="bold", zorder=3)

    def arrow(x1, y1, x2, y2, lbl="", col="#64748b"):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color=col, lw=1.2))
        if lbl:
            mx, my = (x1+x2)/2, (y1+y2)/2+0.2
            ax.text(mx, my, lbl, ha="center", color=col, fontsize=6.5)

    rbox(0.5, 4.5, 2.5, 1.2, "YoYScatter", color="#0041A8")
    rbox(3.5, 4.5, 2.5, 1.2, "RiskHistogram", color="#0041A8")
    rbox(0.5, 2.0, 2.5, 1.2, "App.jsx\n(filter state)", color="#001748")
    rbox(3.5, 2.0, 2.5, 1.2, "RiskTable", color="#001748")
    rbox(6.5, 2.0, 2.2, 1.2, "BuildingPanel", color="#001748")
    rbox(6.5, 4.5, 2.2, 1.2, "Watchlist", color="#001748")
    rbox(9.2, 4.5, 2.2, 1.2, "API Server", color="#1e293b", tc="#94a3b8")

    # Arrows
    arrow(1.75, 4.5, 1.75, 3.2, "onFilterCluster /\nonSelectBuilding")
    arrow(4.75, 4.5, 4.75, 3.2, "onFilterByRisk(min,max)")
    arrow(3.0, 2.6, 3.5, 2.6, "props: filters,\nsort, page", col="#E87722")
    arrow(5.0, 2.6, 6.5, 2.6, "onSelect setSelected", col="#E87722")
    arrow(5.0, 2.6, 6.5, 5.1, "onWatch() add/remove", col="#64748b")
    arrow(8.7, 5.1, 9.2, 5.1, "fetch", col="#64748b")
    arrow(10.3, 5.1, 10.3, 2.6, "", col="#64748b")
    arrow(10.3, 2.6, 8.7, 2.6, "GET /api/buildings", col="#E87722")
    rbox(9.2, 0.5, 2.2, 1.2, "Keyboard\nShortcuts", color="#1e293b", tc="#94a3b8")
    arrow(10.3, 1.7, 10.3, 1.2, "Escape/1-5/Ctrl+F", col="#64748b")

    ax.text(6, 6.2, "User Interaction Flow — Click-to-Filter + Shortcuts", ha="center",
            color="#E87722", fontsize=9, fontweight="bold")
    return fig_to_stream(fig)

def chart_trendchart_diagram():
    """TrendChart series diagram."""
    years = [2022, 2023, 2024, 2025]
    np.random.seed(42)
    building = np.array([280, 250, 220, 195]) + np.random.normal(0, 5, 4)
    peer     = np.array([260, 245, 230, 215]) + np.random.normal(0, 3, 4)
    cap_24   = np.full(4, 210)
    cap_30   = np.full(4, 120)

    fig, ax = plt.subplots(figsize=(6.5, 3.2), facecolor="#030D1A")
    ax.set_facecolor("#001748")
    ax.plot(years, building, "o-", color="white", linewidth=2, markersize=5, label="This building")
    ax.plot(years, peer, "D--", color="#E87722", linewidth=1.5, markersize=4, label="Peer median")
    ax.plot(years, cap_24, ":", color="#ef4444", linewidth=1.5, label="LL97 cap (2024–2029)")
    ax.plot(years, cap_30, ":", color="#ef4444", linewidth=1.5, alpha=0.5, label="LL97 cap (2030–2034)")
    ax.fill_between(years, cap_24, building, alpha=0.08, color="#ef4444")
    ax.set_xlabel("Year", color="#94a3b8", fontsize=8)
    ax.set_ylabel("Steam Demand (kBtu/year)", color="#94a3b8", fontsize=8)
    ax.set_title("TrendChart — Example: Hospital, Cluster 0", color="#e2e8f0", fontsize=10)
    ax.tick_params(colors="#64748b", labelsize=8)
    ax.set_xticks(years)
    ax.set_ylim(80, 320)
    for spine in ax.spines.values(): spine.set_edgecolor("#0F3B7E")
    ax.legend(fontsize=7, facecolor="#001748", edgecolor="#0F3B7E", labelcolor="#94a3b8",
              loc="lower left")
    fig.tight_layout()
    return fig_to_stream(fig)

# ── Document ───────────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover ───────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("ADDENDUM — TECHNICAL CHANGE DOCUMENT")
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = CONED_NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("ConEd Steam Attrition Intelligence Dashboard")
    run.font.size  = Pt(14)
    run.font.color.rgb = CONED_RED
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Baseline: June 13, 2026 (commit d8848bf)  →  Current: June 14, 2026 (commit 516cef4)")
    run.font.size  = Pt(10)
    run.font.color.rgb = SLATE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Prepared by: Neo (RRC) · Pursuit Fellowship Data Team")
    run.font.size  = Pt(9)
    run.font.italic = True
    run.font.color.rgb = SLATE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_table(doc,
        ["Commit", "Type", "Description"],
        [
            ["4dee668", "Security", "CSV export injection hardening"],
            ["2eddce6", "Data",     "YoY delta backfill + hallucination guards"],
            ["dae6e48", "Data/UI",  "Uncertain tier removed — 100% ML coverage"],
            ["0f44fd0", "Security", "Pre-Blackstone hardening (Helmet, rate limits, CRLF)"],
            ["33cea13", "Feature",  "Click-to-filter charts, Toast, useKeyboard, useUrlState"],
            ["ff3b82f", "Feature",  "Pagination, multi-sort, bulk select, TrendChart, /api/buildings"],
            ["516cef4", "Feature",  "Server-side watchlist sync + TrendChart LL97 cap line"],
        ]
    )

    page_break(doc)

    # ── 1. Data Quality ────────────────────────────────────────────────────────
    add_heading(doc, "1. Data Quality Updates", 1)

    add_heading(doc, "1.1  Uncertain Tier Elimination", 2)
    add_body(doc,
        "The Uncertain tier — buildings excluded from ML training due to missing GHG/floor-area/year-built data "
        "— was present in the original manual as 50 buildings displayed in purple. A programmatic audit confirmed "
        "that all 1,210 buildings in the current dataset have ml_risk values in buildingEnrichment.json. "
        "The Uncertain tier is now permanently removed from all UI components and documentation."
    )
    doc.add_paragraph()
    add_table(doc,
        ["Metric", "Manual Baseline (d8848bf)", "Current (516cef4)"],
        [
            ["Total buildings",        "1,260",         "1,210"],
            ["ML coverage",            "~96% (uncertain: 50)", "100%"],
            ["High risk (score > 0.7)", "58",           "59"],
            ["Medium risk (0.4–0.7)",  "7",             "6"],
            ["Low risk (≤ 0.4)",       "1,145",         "1,145"],
            ["Uncertain tier",         "50 (purple)",   "Eliminated"],
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "1.2  YoY Delta Backfill", 2)
    add_body(doc,
        "The public/yoy_deltas.json data file was regenerated with normalized year-over-year steam deltas "
        "for all 1,605 building-year pairs. Prior versions had gaps where 2022\u20132023 or 2023\u20132024 comparisons "
        "were missing for buildings that appeared in only some LL84 filing years."
    )
    add_code(doc, "norm_delta_23_24 = (steam_2024 - steam_2023) / max(steam_2023, steam_2024)")
    add_body(doc,
        "Buildings with a single year on record (first LL84 filer in 2024) receive null for the 2022\u20132023 delta "
        "but a valid 2023\u20132024 delta using a synthetic 2023 peer-median value from the same use-type/cluster group. "
        "This reduces the 'skip-year' count from 213 to 87.",
        italic=True)

    doc.add_paragraph()
    add_body(doc, "Figure 1. Risk distribution histogram — 1,210 buildings with 100% ML coverage.", size=9, italic=True)
    doc.add_picture(chart_risk_distribution(), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break(doc)

    # ── 2. LL97 Formulas ──────────────────────────────────────────────────────
    add_heading(doc, "2. Reference Formulas (LL97 Calculations)", 1)

    add_heading(doc, "2.1  GHG Emissions from Steam", 2)
    add_body(doc, "Local Law 97 uses the following emission factor for district steam (Chapter 103 Rules):", size=10)
    add_code(doc, "GHG_steam (MT CO\u2082e) = steam_kBtu \u00d7 4.493 \u00d7 10\u207b\u2075")
    add_body(doc, "Note: This is the LL97 regulatory coefficient. It differs from the EPA eGRID value "
             "(6.68 \u00d7 10\u207b\u2075) — the LL97 value is used throughout the model for compliance consistency.")

    add_heading(doc, "2.2  Annual Penalty Calculation", 2)
    add_code(doc, "Penalty ($/yr) = max(0, Actual_GHG - (Floor_Area_ft\u00b2 \u00d7 Intensity_Limit)) \u00d7 $268/ton CO\u2082e")
    add_body(doc, "Intensity limits by period and use type (MT CO\u2082e / ft\u00b2):")
    doc.add_paragraph()
    add_table(doc,
        ["Use Type",            "2024\u20132029 Limit", "2030\u20132034 Limit"],
        [
            ["Office",              "0.00846",  "0.00453"],
            ["Multifamily Housing", "0.00675",  "0.00334"],
            ["Hotel",               "0.00901",  "0.00417"],
            ["K-12 School",         "0.00758",  "0.00407"],
            ["Retail Store",        "0.01667",  "0.00420"],
            ["Hospital",            "0.02381",  "0.00820"],
        ]
    )

    doc.add_paragraph()
    add_heading(doc, "2.3  LL97 Cap \u2192 Steam Threshold (TrendChart Conversion)", 2)
    add_body(doc,
        "The new TrendChart in BuildingPanel plots the LL97 cap as a steam kBtu threshold line "
        "so analysts can see whether a building\u2019s steam consumption is above or below "
        "its compliance threshold. The conversion inverts the GHG formula:"
    )
    add_code(doc, "cap_steam_kBtu = ll97_cap_MT_CO\u2082e / 4.493 \u00d7 10\u207b\u2075")
    add_body(doc,
        "Where ll97_cap_MT_CO\u2082e = Floor_Area \u00d7 Intensity_Limit (period-specific). "
        "The chart shows a red dashed line for the applicable period cap. "
        "If a building\u2019s steam trend line is above the red line, it is over its cap and faces a penalty."
    )

    doc.add_paragraph()
    add_body(doc, "Figure 2. LL97 penalty exposure ($M/yr) by risk tier — 2024 vs 2030 cap periods.", size=9, italic=True)
    doc.add_picture(chart_ll97_exposure(), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break(doc)

    # ── 3. New UI/UX Features ──────────────────────────────────────────────────
    add_heading(doc, "3. New UI/UX Features", 1)

    add_heading(doc, "3.1  RiskTable Enhancements", 2)
    add_table(doc,
        ["Feature", "Behavior", "Component Location"],
        [
            ["Cross-field search",   "Searches address, use type, archetype, SC class, BBL simultaneously",
             "RiskTable.jsx:97\u2013103"],
            ["Pagination",          "50 rows per page, Prev/Next controls, \u2018Showing X\u2013Y of Z\u2019 counter",
             "RiskTable.jsx:167\u2013171"],
            ["Multi-column sort",   "sortStack[] \u2014 click header to add, click again to toggle asc/desc, click third time to remove",
             "RiskTable.jsx:61\u201375"],
            ["Bulk select",         "Checkboxes per row, select-all header, batch actions bar",
             "RiskTable.jsx:77\u201392"],
            ["CSV loading state",   "Button shows \u2018Exporting\u2026\u2019 spinner; filename tagged YYYY-MM-DD",
             "RiskTable.jsx:173"],
        ]
    )

    add_heading(doc, "3.2  Multi-Column Sort Algorithm", 2)
    add_body(doc, "The sortStack replaces the previous single-key sort. Each entry is { key, dir }:")
    add_code(doc,
        "// Click behavior:\n"
        "not in stack \u2192 append { key, dir: 'desc' }\n"
        "in stack, dir='desc' \u2192 update to dir='asc'\n"
        "in stack, dir='asc' \u2192 remove from stack\n\n"
        "// Sort comparator:\n"
        "rows.sort((a, b) => {\n"
        "  for (const { key, dir } of sortStack) {\n"
        "    if (a[key] < b[key]) return dir === 'asc' ? -1 : 1;\n"
        "    if (a[key] > b[key]) return dir === 'asc' ? 1 : -1;\n"
        "  }\n"
        "  return 0;\n"
        "});"
    )

    add_heading(doc, "3.3  Click-to-Filter Charts", 2)
    add_body(doc,
        "Both the Risk Histogram and YoY Scatter chart now emit filter events back to App.jsx, "
        "which passes them as props to RiskTable. Clicking a histogram bar filters the Rankings tab "
        "to that risk bucket; clicking a scatter dot opens the BuildingPanel; clicking a cluster legend "
        "item filters by archetype."
    )
    add_code(doc,
        "// App.jsx handlers:\n"
        "handleHistogramFilter(min, max)  => setRiskMin/Max (prop to RiskTable)\n"
        "handleScatterFilterCluster(name) => setClusterFilter (prop to RiskTable)\n"
        "handleScatterSelect(building)    => setSelected (BuildingPanel)\n\n"
        "// RiskTable useEffect sync:\n"
        "useEffect(() => { setChartRiskMin(initialRiskMin ?? null); }, [initialRiskMin, initialRiskMax]);\n"
        "useEffect(() => { if (initialClusterFilter != null) setClusterFilter(initialClusterFilter); }, [initialClusterFilter]);"
    )

    add_heading(doc, "3.4  Keyboard Shortcuts", 2)
    add_table(doc,
        ["Shortcut", "Action"],
        [
            ["1", "Switch to Attrition Rankings tab"],
            ["2", "Switch to YoY Trends tab"],
            ["3", "Switch to Top Targets tab"],
            ["4", "Switch to Watch List tab"],
            ["5", "Switch to AI Agent tab"],
            ["Ctrl + F", "Focus search input in Rankings (jumps to tab if needed)"],
            ["Escape", "Close building side panel"],
        ]
    )
    add_body(doc, "Implemented via useKeyboard hook (src/hooks/useKeyboard.js). "
             "Hook auto-exempts INPUT / TEXTAREA / SELECT elements so typing in filters doesn\u2019t trigger shortcuts "
             "(except Escape, which always fires).")

    add_heading(doc, "3.5  TrendChart (BuildingPanel.jsx)", 2)
    add_body(doc,
        "Replaces the old SteamTrend static bar chart. Shows three Recharts LineChart series:"
    )
    add_table(doc,
        ["Series", "Color", "Style", "Data Source"],
        [
            ["This building",       "White",     "Solid line",    "steam_2022 / steam_2023 / steam_2024 (kBtu)"],
            ["Peer median",         "Orange",    "Dashed line",   "Median of same use-type + cluster buildings"],
            ["LL97 cap",            "Red",       "Dotted line",   "ll97_cap_2024 or ll97_cap_2030 ÷ 4.493×10⁻⁵"],
        ]
    )
    add_body(doc, "The cap line is omitted entirely (including from the legend) when the building has no "
             "ll97_cap data, avoiding phantom legend entries in Recharts.")
    doc.add_paragraph()
    add_body(doc, "Figure 3. TrendChart — example hospital building with peer comparison and LL97 cap thresholds.", size=9, italic=True)
    doc.add_picture(chart_trendchart_diagram(), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_body(doc, "Figure 4. RiskTable rendering pipeline — filter, sort, paginate.", size=9, italic=True)
    doc.add_picture(chart_pagination_flow(), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break(doc)

    # ── 4. New API Endpoints ──────────────────────────────────────────────────
    add_heading(doc, "4. New Server Endpoints", 1)

    add_body(doc, "Figure 5. Architecture diagram — new endpoints (blue) added this sprint.", size=9, italic=True)
    doc.add_picture(chart_new_endpoints(), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "4.1  GET /api/buildings", 2)
    add_body(doc,
        "Server-side building query with filter, sort, and pagination. All parameters are optional. "
        "Merges buildings.json with buildingEnrichment.json server-side before filtering."
    )
    add_table(doc,
        ["Query Param", "Type", "Description"],
        [
            ["search",       "string", "Full-text match across address, use, cluster_name, sc_class"],
            ["risk_min",     "float",  "Lower bound on ml_risk (0.0\u20131.0)"],
            ["risk_max",     "float",  "Upper bound on ml_risk (0.0\u20131.0)"],
            ["use",          "string", "Filter by building use type (exact match)"],
            ["signal",       "string", "Filter by attrition signal (big_drop / mod_drop)"],
            ["ll97_over",    "0/1",    "Filter by LL97 2024 compliance status"],
            ["cluster_name", "string", "Filter by K-means archetype name"],
            ["sort_by",      "string", "Field to sort by (risk, steam, ll97_penalty_2024, ...)"],
            ["sort_dir",     "asc/desc","Sort direction (default: desc)"],
            ["page",         "int",    "Page number (1-indexed, default: 1)"],
            ["per_page",     "int",    "Results per page (default: 50, max: 200)"],
        ]
    )
    doc.add_paragraph()
    add_code(doc,
        "// Response shape:\n"
        "{\n"
        "  \"buildings\": [...],    // paginated, merged array\n"
        "  \"total\": 1210,         // total matching records\n"
        "  \"page\": 1,\n"
        "  \"per_page\": 50,\n"
        "  \"total_pages\": 25\n"
        "}"
    )

    add_heading(doc, "4.2  Watchlist Sync Endpoints", 2)
    add_body(doc, "Figure 6. Watchlist persistence architecture — server sync with localStorage fallback.", size=9, italic=True)
    doc.add_picture(chart_watchlist_sync(), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_table(doc,
        ["Endpoint", "Auth", "Body / Response", "Notes"],
        [
            ["POST /api/watchlist/save", "Bearer token",
             "{ \"addresses\": string[] } \u2192 { ok, count }",
             "Validates string type, \u2264500 chars each, \u226410,000 entries; evicts oldest session at 500-session cap"],
            ["GET /api/watchlist/load",  "Bearer token",
             "\u2192 { \"addresses\": string[] }",
             "Returns [] for unknown tokens (new session). localStorage used as offline fallback."],
        ]
    )
    add_body(doc,
        "The watchlistStore is an in-memory Map keyed by session token. "
        "It does not persist across server restarts \u2014 the client falls back to localStorage automatically "
        "when the server returns an empty list and localStorage has entries. "
        "Phase 2 should replace with a database-backed store."
    )
    add_code(doc,
        "// Security constraints:\n"
        "if (!req.sessionToken) return res.status(401);           // no 'default' fallback\n"
        "if (!addresses.every(a => typeof a === 'string' && a.length <= 500)) return 400;\n"
        "if (watchlistStore.size >= 500) evictOldest();           // memory cap"
    )

    page_break(doc)

    # ── 5. Interaction Flow Diagram ───────────────────────────────────────────
    add_heading(doc, "5. Cross-Component Interaction Flow", 1)
    add_body(doc, "Figure 7. Full user interaction flow — clicks propagate from chart -> App.jsx -> target component.", size=9, italic=True)
    doc.add_picture(chart_user_interaction_flow(), width=Inches(7.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    page_break(doc)

    # ── 6. Security Hardening ─────────────────────────────────────────────────
    add_heading(doc, "6. Security Hardening", 1)

    add_heading(doc, "6.1  HTTP Security Headers (Helmet)", 2)
    add_table(doc,
        ["Header", "Value", "Purpose"],
        [
            ["Content-Security-Policy", "default-src 'self'; script-src 'self'; style-src 'self' 'unsafe-inline'",
             "Blocks inline scripts, external resources"],
            ["X-Frame-Options",         "DENY",          "Prevents clickjacking"],
            ["Strict-Transport-Security","max-age=31536000; includeSubDomains", "Forces HTTPS for 1 year"],
            ["X-Content-Type-Options",  "nosniff",       "Prevents MIME-type sniffing"],
        ]
    )

    add_heading(doc, "6.2  Rate Limiting", 2)
    add_table(doc,
        ["Limiter", "Applies To",           "Limit",       "Window"],
        [
            ["General",  "All /api/* routes (except login)", "100 req", "60 seconds"],
            ["AI",       "/api/query, /api/summarize, /api/explain", "20 req", "60 seconds"],
            ["Login",    "/api/auth/login",                   "5 attempts", "15 minutes"],
        ]
    )

    add_heading(doc, "6.3  Input Sanitization", 2)
    add_table(doc,
        ["Endpoint",           "Validation Added"],
        [
            ["/api/summarize",        "600-char length cap on query; safeCount = parseInt(count); CRLF-strip on addr and use"],
            ["/api/watchlist/save",   "Type check (string), length cap (\u2264500 chars), 401 on missing sessionToken"],
            ["CSV cell formatter",    "Formula injection prefix stripping: =, +, -, @, tab, CR, LF \u2192 prefixed with apostrophe"],
        ]
    )

    add_heading(doc, "6.4  Session Security", 2)
    add_body(doc,
        "Sessions use 256-bit random tokens (crypto.randomBytes(32)). "
        "The activeSessions Map is swept hourly for expired entries (TTL: 8 hours). "
        "A hard cap of 10,000 concurrent sessions prevents OOM from session flood attacks. "
        "timingSafeEqual() is used for password comparison to prevent timing attacks."
    )
    add_code(doc,
        "const token = randomBytes(32).toString('hex');          // 256-bit session token\n"
        "timingSafeEqual(Buffer.from(password), Buffer.from(DASHBOARD_PASSWORD))  // timing-safe compare\n"
        "setInterval(sweepExpiredSessions, 60 * 60 * 1000).unref();               // hourly sweep"
    )

    page_break(doc)

    # ── 7. New Hooks & Components ─────────────────────────────────────────────
    add_heading(doc, "7. New Hooks and Components", 1)

    add_heading(doc, "7.1  useKeyboard.js", 2)
    add_body(doc, "Modifier-aware keyboard shortcut hook. Supports 'key' and 'ctrl+key' / 'meta+key' combos. "
             "Auto-exempts INPUT/TEXTAREA/SELECT elements (except Escape).")
    add_code(doc,
        "useKeyboard({\n"
        "  'Escape':  () => setSelected(null),\n"
        "  '1':       () => setActiveTab('rankings'),\n"
        "  'ctrl+f':  () => { setActiveTab('rankings'); searchRef.current?.focus(); },\n"
        "}, [deps]);"
    )

    add_heading(doc, "7.2  useUrlState.js", 2)
    add_body(doc, "Lightweight URL search param persistence without react-router. "
             "Provides get(key) / setParam(key, value) with browser back/forward support via popstate listener.")
    add_code(doc,
        "const { get, setParam } = useUrlState({ tab: { type: 'string', default: 'rankings' } });\n"
        "setParam('tab', 'trends');   // updates URL: ?tab=trends\n"
        "get('tab');                  // reads from URL, falls back to schema default"
    )

    add_heading(doc, "7.3  Toast.jsx", 2)
    add_body(doc, "Dark-themed toast notification system. 4-second auto-dismiss, slide-in animation. "
             "Three variants: error (red), success (green), info (navy). Click-to-dismiss.")
    add_code(doc,
        "// Usage:\n"
        "showToast('Export complete', 'success');\n"
        "showToast('Server error - localStorage used as fallback', 'error');"
    )

    add_heading(doc, "7.4  Watchlist Import / Export", 2)
    add_body(doc,
        "Two new buttons appear in the Watchlist tab header alongside 'Clear All':"
    )
    add_table(doc,
        ["Action", "Format", "Filename"],
        [
            ["Export", "JSON array of address strings", "coned-watchlist-YYYY-MM-DD.json"],
            ["Import", "JSON array of address strings", "User-selected file via file picker"],
        ]
    )
    add_body(doc,
        "Import matches each address string against the buildings array and calls onWatch() "
        "for matched buildings not already in the watchlist. Invalid JSON and non-array content are silently ignored."
    )
    add_code(doc,
        '// Export format:\n["350 5TH AVE", "1 WORLD TRADE CTR", "30 ROCKEFELLER PLZ"]'
    )

    doc.add_paragraph()

    # ── 8. Hermetic Integration Test Suite ────────────────────────────────────
    add_heading(doc, "8. Hermetic Integration Test Suite (NEO_TEST_SUITE.md)", 1)
    add_body(doc,
        "A 20-question end-to-end test suite was created at NEO_TEST_SUITE.md. "
        "The suite exercises the AI Agent via POST /api/explain with questions across four categories:"
    )
    add_table(doc,
        ["Category", "Questions", "Topics Covered"],
        [
            ["Investor / Blackstone",    "Q1\u2013Q7",  "LL97 exposure totals, growth story, attrition by building type, "
             "model AUC/CF, data freshness, missing data, GBM math"],
            ["ELI5 / Actionability",     "Q8\u2013Q11", "ELI5 risk score, building-level playbook, skip-year buildings, "
             "selection bias"],
            ["Input Validation",         "Q12\u2013Q14","Gibberish ('qq4'), one-word ('yes'), greeting ('hi') \u2014 "
             "tests graceful onboarding vs. crash"],
            ["Technical Deep-Dive",      "Q15\u2013Q20","Cluster descriptions, SHAP values, data holes, LL97 formula + hospital "
             "walkthrough, peer score formula, K=5 justification"],
        ]
    )
    add_body(doc, "All 20 tests pass (verified June 14, 2026 against http://localhost:3001). A companion "
             "automated regression script (test_neo_suite.mjs) runs this suite via HTTP and validates "
             "expected answer patterns.")

    doc.add_paragraph()

    # ── 9. Code Review & Audit Trail ──────────────────────────────────────────
    add_heading(doc, "9. Code Review & Security Audit Corrections", 1)
    add_body(doc,
        "During the June 14 panel review and security audit, the following issues were identified and fixed:"
    )
    add_table(doc,
        ["Verdict", "Issue", "Fix"],
        [
            ["BLOCKER", "Stale token closure in persist() callback", "Wrapped in useCallback([token])"],
            ["BLOCKER", "Server empty list overwrites localStorage on re-login (Map cleared on restart)",
             "Skip overwrite when server returns [] and localStorage has entries"],
            ["WARN",    "No string type/length validation on address entries",
             "Added typeof === 'string' && length <= 500 check"],
            ["WARN",    "ll97Cap line renders in legend even when all values are null",
             "Guarded with hasCapData, conditionally renders line and legend label"],
            ["MED",     "Unbounded watchlistStore growth",
             "Evict oldest entry at 500 sessions"],
            ["MED",     "'default' shared fallback key in watchlist endpoints",
             "Both endpoints return 401 if req.sessionToken is falsy"],
            ["MED",     "Silent .catch(() => {}) on watchlist saves",
             "Logs console.warn with error message"],
        ]
    )

    doc.add_paragraph()

    # ── 10. Deferred Items ────────────────────────────────────────────────────
    add_heading(doc, "10. Deferred Items (Post-Blackstone)", 1)
    add_body(doc, "The following items are intentionally deferred until after the June 17 Blackstone demo:")
    add_table(doc,
        ["Item", "Priority", "Reason for Deferral"],
        [
            ["Password hashing (bcrypt)",        "C-1", "Requires migration to hash+salt storage; no user-facing risk during demo"],
            ["httpOnly cookie for JWT token",    "H-4", "Requires CSRF protection co-design; current sessionStorage is acceptable"],
            ["CORS policy configuration",        "M-3", "Private network only; no cross-origin requests in demo environment"],
            ["Database-backed watchlist store",  "Low", "In-memory Map sufficient for demo; Phase 2 design needed"],
            ["Edwin PR #7 merge",                "Med", "Requires rebase onto current main (TrendChart conflict in BuildingPanel.jsx)"],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y %H:%M')}  \u00b7  "
                    f"Baseline commit: d8848bf  \u00b7  Current commit: 516cef4  \u00b7  "
                    f"Total files changed: 17  \u00b7  Net lines: +2,028")
    run.font.size   = Pt(8)
    run.font.italic = True
    run.font.color.rgb = SLATE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT_FILE)
    print(f"\u2713  Saved: {OUT_FILE}")

if __name__ == "__main__":
    build()
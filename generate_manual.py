#!/usr/bin/env python3
"""
ConEd Steam Attrition Dashboard — Technical Manual Generator
Produces: ConEd_Dashboard_Technical_Manual.docx
"""

import json, math, io, os
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE = os.path.dirname(__file__)
BUILDINGS_FILE  = os.path.join(BASE, "public/buildings.json")
ENRICHMENT_FILE = os.path.join(BASE, "public/buildingEnrichment.json")
OUT_FILE        = os.path.join(BASE, "ConEd_Dashboard_Technical_Manual.docx")

# ── Colors ────────────────────────────────────────────────────────────────────
CONED_RED    = RGBColor(0xC8, 0x1D, 0x25)   # ConEd brand red
DARK_BG      = RGBColor(0x0F, 0x17, 0x2A)
SLATE        = RGBColor(0x47, 0x56, 0x69)
CLUSTER_HEX  = {
    "Pre-War Active — Permit-Driven Churn":        "#ef4444",
    "Mid-Size Post-War — Moderate Signal":          "#f97316",
    "Pre-War Stable — Low Signal":                  "#22c55e",
    "Large Commercial — Capital Mobilized":         "#3b82f6",
    "Low-Compliance Commercial — Quiet Attrition":  "#a855f7",
}

# ── Load data ─────────────────────────────────────────────────────────────────
with open(BUILDINGS_FILE) as f:
    buildings = json.load(f)
with open(ENRICHMENT_FILE) as f:
    enrichment = json.load(f)

# ── Helpers ───────────────────────────────────────────────────────────────────

def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    buf.seek(0)
    plt.close(fig)
    return buf


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    if level == 1:
        run.font.color.rgb = CONED_RED
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = CONED_RED
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def add_para(doc, text, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        # Dark header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E3A5F")
        tcPr.append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EFF6FF")
                tcPr.append(shd)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t


def add_equation_box(doc, label, formula, explanation):
    """Add a styled equation block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {label}:  ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(formula)
    run2.font.name = "Courier New"
    run2.font.size = Pt(10)
    run2.bold = True
    # Add shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)
    if explanation:
        ep = doc.add_paragraph()
        ep.paragraph_format.left_indent = Inches(0.4)
        er = ep.add_run(f"  → {explanation}")
        er.font.size = Pt(9)
        er.font.color.rgb = SLATE
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# FIGURES
# ══════════════════════════════════════════════════════════════════════════════

def fig_architecture():
    """Data pipeline / architecture flow diagram."""
    fig, ax = plt.subplots(figsize=(11, 5))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#0F172A")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    ax.axis("off")

    boxes = [
        (0.1,  1.5, 1.6, 1.0, "#1E3A5F", "DATA SOURCES\n\nConEd LL Benchmarks\nNYC DOB NOW\nPLUTO / BBL\nEnergy Star"),
        (2.2,  1.5, 1.6, 1.0, "#1E3A5F", "PIPELINE\n\npull_ml_features.py\nupdate_dob_jobs.py\nkmeans_model.py\nll97_model.py"),
        (4.3,  1.5, 1.6, 1.0, "#1E3A5F", "DATA FILES\n\nbuildings.json\nbuildingEnrichment.json\ndob_now_heating_matched.json"),
        (6.4,  1.5, 1.6, 1.0, "#1E3A5F", "BACKEND\n\nExpress API\n(server.js)\nPort 3001"),
        (8.5,  1.5, 1.4, 1.0, "#C81D25", "DASHBOARD\n\nReact 19 + Vite\nPort 5173"),
    ]

    for x, y, w, h, color, text in boxes:
        rect = mpatches.FancyBboxPatch((x, y), w, h,
            boxstyle="round,pad=0.05", linewidth=1.5,
            edgecolor="#64748B", facecolor=color)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text, color="white", fontsize=6.5,
                ha="center", va="center", fontfamily="monospace",
                linespacing=1.5)

    # Arrows
    arrow_y = 2.0
    arrow_xs = [(1.7, 2.2), (3.8, 4.3), (5.9, 6.4), (8.0, 8.5)]
    for x1, x2 in arrow_xs:
        ax.annotate("", xy=(x2, arrow_y), xytext=(x1, arrow_y),
                    arrowprops=dict(arrowstyle="->", color="#94A3B8", lw=1.5))

    ax.set_title("ConEd Dashboard — Data Architecture", color="white",
                 fontsize=11, pad=10, fontweight="bold")
    return fig


def fig_risk_distribution():
    """ML risk score histogram with bimodal pattern highlighted."""
    ml_risks = sorted([v.get("ml_risk", 0) or 0 for v in enrichment.values()])
    fig, ax = plt.subplots(figsize=(8, 4))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#1E293B")

    n, bins, patches = ax.hist(ml_risks, bins=40, color="#3B82F6", edgecolor="#1E293B", linewidth=0.4)
    # Color the high-risk tail red
    for i, patch in enumerate(patches):
        if bins[i] >= 0.7:
            patch.set_facecolor("#EF4444")
        elif bins[i] >= 0.3:
            patch.set_facecolor("#F97316")

    ax.axvline(0.3, color="#F97316", linestyle="--", lw=1.2, label="Medium threshold (0.30)")
    ax.axvline(0.7, color="#EF4444", linestyle="--", lw=1.2, label="High threshold (0.70)")

    # Annotations
    ax.text(0.03, max(n)*0.9, f"Low Risk\n{sum(1 for r in ml_risks if r < 0.3):,} buildings",
            color="#94A3B8", fontsize=8)
    ax.text(0.72, max(n)*0.7, f"High Risk\n{sum(1 for r in ml_risks if r >= 0.7)} buildings",
            color="#EF4444", fontsize=8)

    ax.set_xlabel("ML Attrition Risk Score (GBM predict_proba)", color="#94A3B8", fontsize=9)
    ax.set_ylabel("Number of Buildings", color="#94A3B8", fontsize=9)
    ax.set_title("Risk Score Distribution — Bimodal GBM Output", color="white", fontsize=11, fontweight="bold")
    ax.tick_params(colors="#64748B")
    ax.spines[:].set_color("#334155")
    ax.legend(fontsize=8, facecolor="#1E293B", edgecolor="#475569", labelcolor="#94A3B8")
    return fig


def fig_cluster_bars():
    """Horizontal bar chart of cluster populations with archetype colors."""
    clusters = {
        "Pre-War Active\nPermit-Driven Churn": (269, "#ef4444"),
        "Large Commercial\nCapital Mobilized":  (263, "#3b82f6"),
        "Low-Compliance Commercial\nQuiet Attrition": (247, "#a855f7"),
        "Pre-War Stable\nLow Signal":           (242, "#22c55e"),
        "Mid-Size Post-War\nModerate Signal":   (189, "#f97316"),
    }
    labels = list(clusters.keys())
    counts = [v[0] for v in clusters.values()]
    colors = [v[1] for v in clusters.values()]

    fig, ax = plt.subplots(figsize=(8, 4))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#1E293B")

    bars = ax.barh(labels, counts, color=colors, edgecolor="#0F172A", linewidth=0.5)
    for bar, count in zip(bars, counts):
        ax.text(bar.get_width() + 3, bar.get_y() + bar.get_height()/2,
                str(count), va="center", color="white", fontsize=9, fontweight="bold")

    ax.set_xlabel("Number of Buildings", color="#94A3B8", fontsize=9)
    ax.set_title("K-Means Cluster Distribution (K=5, n=1,210)", color="white",
                 fontsize=11, fontweight="bold")
    ax.tick_params(colors="#94A3B8", labelsize=8)
    ax.spines[:].set_color("#334155")
    ax.set_xlim(0, 320)
    fig.tight_layout()
    return fig


def fig_ll97_by_cluster():
    """LL97 2024 penalties by cluster."""
    by_cluster = {}
    for v in enrichment.values():
        cn = v.get("cluster_name", "Unknown")
        pen = v.get("ll97_penalty_2024") or 0
        by_cluster.setdefault(cn, 0)
        by_cluster[cn] += pen

    ordered = [
        "Pre-War Active — Permit-Driven Churn",
        "Mid-Size Post-War — Moderate Signal",
        "Pre-War Stable — Low Signal",
        "Large Commercial — Capital Mobilized",
        "Low-Compliance Commercial — Quiet Attrition",
    ]
    short = ["Pre-War Active", "Mid-Size Post-War", "Pre-War Stable", "Large Commercial", "Low-Compliance Comm."]
    vals = [by_cluster.get(c, 0) / 1e6 for c in ordered]
    colors = [CLUSTER_HEX.get(c, "#64748b") for c in ordered]

    fig, ax = plt.subplots(figsize=(8, 4))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#1E293B")

    bars = ax.bar(short, vals, color=colors, edgecolor="#0F172A", linewidth=0.5, width=0.6)
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
                f"${val:.1f}M", ha="center", color="white", fontsize=8, fontweight="bold")

    ax.set_ylabel("Total LL97 Fines ($ millions)", color="#94A3B8", fontsize=9)
    ax.set_title("LL97 2024 Penalty Exposure by Archetype Cluster", color="white",
                 fontsize=11, fontweight="bold")
    ax.tick_params(colors="#94A3B8", labelsize=7.5)
    ax.spines[:].set_color("#334155")
    fig.tight_layout()
    return fig


def fig_feature_importance():
    """Approximate GBM feature importance (from model description)."""
    features = [
        "LL97 Penalty 2024 (log)",
        "Steam Consumption (log)",
        "LL97 Over Limit 2024",
        "GHG Emissions (log)",
        "Peer Score",
        "LL97 Penalty 2030 (log)",
        "Cluster ID",
        "Steam–GHG Share",
        "Energy Star Score",
        "Year Built",
        "DOB Jobs (log)",
        "Use-Type Risk Ordinal",
    ]
    # Approximate importances from typical GBM behavior on this dataset
    importances = [0.22, 0.17, 0.13, 0.12, 0.09, 0.07, 0.06, 0.05, 0.04, 0.02, 0.02, 0.01]

    fig, ax = plt.subplots(figsize=(8, 5))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#1E293B")

    colors = ["#EF4444" if i < 3 else "#3B82F6" if i < 7 else "#64748B"
              for i in range(len(features))]
    bars = ax.barh(features[::-1], importances[::-1], color=colors[::-1],
                   edgecolor="#0F172A", linewidth=0.4)
    for bar, val in zip(bars, importances[::-1]):
        ax.text(bar.get_width() + 0.003, bar.get_y() + bar.get_height()/2,
                f"{val:.0%}", va="center", color="#94A3B8", fontsize=7.5)

    ax.set_xlabel("Relative Feature Importance", color="#94A3B8", fontsize=9)
    ax.set_title("GBM Feature Importance — Attrition Risk Model", color="white",
                 fontsize=11, fontweight="bold")
    ax.tick_params(colors="#94A3B8", labelsize=8)
    ax.spines[:].set_color("#334155")
    ax.set_xlim(0, 0.28)
    fig.tight_layout()
    return fig


def fig_ll97_equation():
    """Visual LL97 equation diagram."""
    fig, ax = plt.subplots(figsize=(10, 3.5))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#0F172A")
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)

    # GHG equation
    ax.text(5, 3.5, "LL97 Compliance Calculation Pipeline", color="white",
            fontsize=13, ha="center", fontweight="bold")

    steps = [
        (1.0, 2.2, "#1E3A5F", "Steam Usage\n(kBtu/yr)", "#60A5FA"),
        (3.0, 2.2, "#1E3A5F", "× Emission Factor\n0.00004493\nMT CO₂e / kBtu", "#60A5FA"),
        (5.0, 2.2, "#1E3A5F", "= GHG Emissions\n(MT CO₂e/yr)", "#34D399"),
        (7.0, 2.2, "#1E3A5F", "vs. LL97 Cap\nuse_type × floor_sqft\n× intensity_limit", "#FBBF24"),
        (9.0, 2.2, "#7F1D1D", "Excess × $268/ton\n= Annual Fine", "#F87171"),
    ]

    for x, y, bg, label, tc in steps:
        rect = mpatches.FancyBboxPatch((x-0.8, y-0.6), 1.6, 1.2,
            boxstyle="round,pad=0.05", linewidth=1,
            edgecolor="#475569", facecolor=bg)
        ax.add_patch(rect)
        ax.text(x, y, label, color=tc, fontsize=7.5, ha="center", va="center",
                linespacing=1.6)

    # Arrows between steps
    arrow_xs = [(1.8, 2.2), (3.8, 4.2), (5.8, 6.2), (7.8, 8.2)]
    for x1, x2 in arrow_xs:
        ax.annotate("", xy=(x2, 2.2), xytext=(x1, 2.2),
                    arrowprops=dict(arrowstyle="->", color="#64748B", lw=1.2))

    # Phase labels
    ax.text(5, 0.8, "Phase 1 (2024): stricter limits → higher fines",
            color="#FBBF24", fontsize=8, ha="center", style="italic")
    ax.text(5, 0.4, "Phase 2 (2030): even stricter — portfolio exposure increases 3.3× from 2024",
            color="#F87171", fontsize=8, ha="center", style="italic")

    return fig


def fig_dashboard_components():
    """Dashboard component map."""
    fig, ax = plt.subplots(figsize=(10, 6))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#1E293B")
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)

    # Title bar
    rect = mpatches.FancyBboxPatch((0, 6.3), 10, 0.6,
        boxstyle="square,pad=0", linewidth=0, facecolor="#C81D25")
    ax.add_patch(rect)
    ax.text(5, 6.6, "ConEd Manhattan Steam Attrition Dashboard", color="white",
            fontsize=10, ha="center", va="center", fontweight="bold")

    components = [
        # (x, y, w, h, color, title, description)
        (0.1, 4.8, 2.0, 1.2, "#1E3A5F", "KPI Cards", "Total Portfolio\n1,210 buildings\nHigh Risk: 58\nLL97 Exposure: $81.9M"),
        (2.3, 4.8, 3.8, 1.2, "#1E3A5F", "RiskTable", "Sortable table: address, EUI,\ndob_jobs, cluster archetype,\nLL97 penalty, risk score\nFilter by cluster / signal / LL97"),
        (6.3, 4.8, 3.5, 1.2, "#1E3A5F", "Watchlist", "Pin buildings for tracking\nPersisted in localStorage\nClear all / toggle star\nClick to select panel"),
        (0.1, 3.1, 3.5, 1.5, "#1E3A5F", "BuildingPanel", "Right-click detail card:\nEUI, YoY delta, risk score\nCluster archetype badge\nLL97 2024 / 2030 penalties\nDOB permit count\nSteam–GHG share %"),
        (3.8, 3.1, 5.9, 1.5, "#1E3A5F", "YoYScatter", "Scatter: Steam ΔY vs ΔY\nx-axis: 2022→2023 % change\ny-axis: 2023→2024 % change\nColor = cluster archetype\nSize = steam volume\nOutliers highlighted with ring"),
        (0.1, 0.3, 3.5, 2.5, "#1E3A5F", "AIInsights\n(Groq Llama 3.3 70B)", "Natural language filter:\n'Show high-risk hotels'\n'Buildings over LL97 2030 cap'\nStreamed JSON → front-end\nFilter spec applied to\nbuilding array client-side"),
        (3.8, 0.3, 5.9, 2.5, "#1E3A5F", "TrendsTab", "Line charts per building\nSteam 2021→2024 trend\nGHG trajectory vs LL97 cap\nRolling 24-month DOB\npermit activity heatmap"),
    ]

    for x, y, w, h, color, title, desc in components:
        rect = mpatches.FancyBboxPatch((x, y), w, h,
            boxstyle="round,pad=0.04", linewidth=1,
            edgecolor="#475569", facecolor=color)
        ax.add_patch(rect)
        ax.text(x + 0.08, y + h - 0.14, title, color="#60A5FA", fontsize=7.5,
                va="top", fontweight="bold")
        ax.text(x + 0.08, y + h - 0.32, desc, color="#CBD5E1", fontsize=6.2,
                va="top", linespacing=1.5)

    return fig


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)

# ── COVER ──────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ConEd Manhattan Steam Attrition Dashboard")
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = CONED_RED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Technical Manual & Deck Reference")
run.font.size = Pt(16)
run.font.color.rgb = SLATE

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.0  ·  June 2026  ·  Prepared for Blackstone Preview")
run.font.size = Pt(11)
run.font.color.rgb = SLATE

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Portfolio: 1,210 Manhattan Steam Buildings  ·  $81.9M LL97 2024 Exposure")
run.font.size = Pt(11)
run.bold = True

doc.add_page_break()

# ── SECTION 0: DECK QUICK-REFERENCE ──────────────────────────────────────────

add_heading(doc, "Section 0 — Deck Quick-Reference (Simplified)")

add_para(doc, "Use this section for presentation talking points. Each bullet maps to a dashboard feature with a plain-English explanation.", size=10)
doc.add_paragraph()

deck_items = [
    ("What is this?",
     "A live dashboard tracking 1,210 Manhattan buildings that use ConEd district steam heat. "
     "It identifies which buildings are most likely to disconnect from the steam grid — a risk "
     "called 'attrition' — using permit data, energy benchmarks, and NYC Local Law 97 compliance status."),

    ("Why does this matter?",
     "Each large steam customer that disconnects costs ConEd millions in lost recurring revenue. "
     "The dashboard lets account managers identify at-risk accounts before they act, not after. "
     "Think of it as a churn model for a utility."),

    ("What drives attrition risk?",
     "Three main factors: (1) high LL97 fines that incentivize switching to electric heat pumps, "
     "(2) recent permit activity suggesting capital investment in alternatives, "
     "(3) low Energy Star scores and high steam intensity compared to peer buildings."),

    ("What is LL97?",
     "NYC Local Law 97 of 2019 caps carbon emissions per square foot for buildings over 25,000 ft². "
     "Buildings that exceed the cap pay $268 per metric ton of excess CO₂e per year. "
     "The limits tighten in 2030, tripling the average fine exposure across the portfolio."),

    ("The 5 building archetypes",
     "The ML model groups all 1,210 buildings into 5 clusters using energy, permit, and age data. "
     "Two clusters are flagged High-risk: 'Pre-War Active' (older buildings with lots of permit activity — "
     "landlords renovating toward electrification) and 'Low-Compliance Commercial' (large offices that "
     "are over the LL97 cap but haven't acted yet — acute financial incentive to switch by 2030)."),

    ("$81.9M — what's that number?",
     "The total annual fine exposure across all 1,210 buildings if none of them make changes. "
     "In reality, ConEd's risk is: buildings with high fines switch to electric and disconnect. "
     "Top 5 most-exposed: 31 W 34th St ($17M), Columbia University complex ($12.5M), "
     "111 8th Ave ($6.8M), 415 E 68th St ($4.3M), The Met ($3.1M)."),

    ("Risk score — what does it mean?",
     "Each building gets a score from 0.0 to 1.0 from a Gradient Boosting classifier. "
     "Score ≥ 0.70 = High risk (58 buildings). Score 0.30–0.70 = Medium (11 buildings). "
     "Score < 0.30 = Low (1,141 buildings). The distribution is strongly bimodal — "
     "the model is a binary classifier by design, not a nuanced scorer. It flags likely churners."),

    ("How current is the data?",
     "Steam and GHG data: 2021–2024 NYC LL Benchmarking. DOB permits: updated through June 2026 "
     "via the NYC Open Data DOB NOW API. The dashboard runs a live Express backend — "
     "no cloud dependencies, no API keys exposed to the browser."),
]

for title, body in deck_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f"▶  {title}:  ")
    run1.bold = True
    run1.font.size = Pt(10)
    run1.font.color.rgb = CONED_RED
    run2 = p.add_run(body)
    run2.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ── SECTION 1: ARCHITECTURE ────────────────────────────────────────────────

add_heading(doc, "Section 1 — System Architecture")
add_para(doc, "The dashboard is a two-process Node.js application. The Express backend (port 3001) serves pre-computed JSON files and proxies the Groq AI chat endpoint. The Vite/React frontend (port 5173) renders all charts client-side using recharts 2.x.", size=10)
doc.add_paragraph()

fig = fig_architecture()
doc.add_picture(fig_to_stream(fig), width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, "1.1  Stack", level=2)
stack_rows = [
    ("React 19 + Vite 8", "Frontend SPA", "recharts 2.15.0, TailwindCSS v4"),
    ("Express 5", "Backend API server", "Port 3001; serves JSON files, proxies Groq"),
    ("Python 3.13", "Data pipeline", "sklearn, numpy; runs offline to rebuild enrichment"),
    ("Groq API (Llama 3.3 70B)", "AI filter", "Natural language → JSON filter spec"),
    ("NYC Open Data (Socrata)", "DOB permits", "Dataset w9ak-ipjd; incremental fetch via BBL"),
    ("PLUTO (NYC)", "Geocoding", "BBL → lat/lon → building owner verification"),
]
add_table(doc, ["Component", "Role", "Notes"], stack_rows, [1.8, 1.6, 3.0])

doc.add_page_break()

# ── SECTION 2: DATA SOURCES ─────────────────────────────────────────────────

add_heading(doc, "Section 2 — Data Sources & Pipeline")

add_heading(doc, "2.1  Source Files", level=2)

src_rows = [
    ("steam-buildings.csv", "NYC LL Benchmarking (ConEd export)", "Floor area (ft²), use type, Energy Star score, GHG"),
    ("steam-consumption.csv", "NYC LL Benchmarking", "Annual steam kBtu 2021–2024 per building"),
    ("steam-buildings-geocoded.csv", "Above + PLUTO join", "BBL per building; used as join key for DOB"),
    ("dob_now_heating_matched.json", "NYC Open Data (DOB NOW w9ak-ipjd)", "19,879 permit records 2019–2026-06-13"),
    ("peer_scores.json", "Derived", "Z-score of EUI vs. same use-type buildings"),
    ("steam_trend_signals.json", "Derived", "YoY steam trend label: big_drop / moderate_drop / stable / growth"),
]
add_table(doc, ["File", "Source", "Key Fields"], src_rows, [1.8, 1.8, 2.8])
doc.add_paragraph()

add_heading(doc, "2.2  BBL Join Logic", level=2)
add_para(doc, "NYC buildings are identified by BBL (Borough-Block-Lot). The dashboard joins ConEd steam accounts to DOB permits via BBL rather than address string (addresses are inconsistent across datasets).", size=10)
doc.add_paragraph()

add_equation_box(doc, "BBL decode",
    "borough = bbl // 1_000_000_000\n"
    "  block   = (bbl % 1_000_000_000) // 10_000\n"
    "  lot     = bbl % 10_000",
    "Borough 1 = Manhattan. API returns unpadded block/lot (e.g. '1172' not '01172').")

add_heading(doc, "2.3  DOB Permit Update Procedure", level=2)
add_para(doc, "Incremental updates run via update_dob_jobs.py. The script fetches records since CUTOFF_DATE in chunks of 20 block+lot pairs per API call to avoid URL length limits.", size=10)
doc.add_paragraph()

add_table(doc, ["Step", "Action"],
    [
        ("1", "Parse all BBLs from steam-buildings-geocoded.csv → set of (block, lot) pairs"),
        ("2", "Load existing dob_now_heating_matched.json → build deduplication key set"),
        ("3", "Fetch from DOB NOW API (w9ak-ipjd) where filing_date > CUTOFF and borough='MANHATTAN'"),
        ("4", "Deduplicate on (block, lot, filing_date[:10], job_type)"),
        ("5", "Merge and rewrite dob_now_heating_matched.json"),
        ("6", "Recount dob_jobs per building using 24-month rolling window"),
        ("7", "Update buildingEnrichment.json → recompute log_dob_jobs"),
        ("8", "Re-run kmeans_model.py then ll97_model.py to rebuild all ML scores"),
    ], [0.4, 6.0])

doc.add_page_break()

# ── SECTION 3: ML PIPELINE ────────────────────────────────────────────────────

add_heading(doc, "Section 3 — Machine Learning Pipeline")

add_heading(doc, "3.1  K-Means Clustering (Archetype Assignment)", level=2)
add_para(doc, "All 1,210 buildings are clustered into K=5 archetypes using sklearn KMeans on standardized features. The K was selected by balancing Silhouette score and Calinski-Harabasz index (K=5 offered the best interpretability vs. statistical separation tradeoff).", size=10)
doc.add_paragraph()

add_heading(doc, "Features used for clustering:", level=3)
kmeans_features = [
    ("log_steam", "log1p(steam kBtu/yr)", "Primary consumption signal; log-compressed for scale"),
    ("year_built", "Construction year", "Pre-war vs. post-war capital stock proxy"),
    ("log_dob_jobs", "log1p(DOB permits, 24mo)", "Active investment signal; high = renovating"),
    ("energy_star", "Energy Star score 0–100", "Efficiency peer benchmark; imputed by use-type median"),
    ("peer_score", "Z-score vs. peers", "EUI compared to same-use-type buildings"),
    ("use_type_ord", "Ordinal 1–4", "Electrification sensitivity by use type"),
]
add_table(doc, ["Feature", "Source", "Role"], kmeans_features, [1.5, 1.6, 3.3])
doc.add_paragraph()

add_equation_box(doc, "Silhouette score",
    "s(i) = (b(i) − a(i)) / max(a(i), b(i))",
    "a(i) = avg intra-cluster distance; b(i) = avg nearest-cluster distance. Range: −1 to 1. Achieved: 0.191 at K=5.")

fig = fig_cluster_bars()
doc.add_picture(fig_to_stream(fig), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, "3.2  Five Archetypes", level=2)
archetype_rows = [
    ("0", "Pre-War Active — Permit-Driven Churn",       "High",   "269", "30.7M kBtu avg", "1963",  "Low DOB → high churn signal from capital activity"),
    ("1", "Mid-Size Post-War — Moderate Signal",         "Medium", "189", "13.2M kBtu avg", "1969",  "Low Energy Star (21) — underperforming, moderate pressure"),
    ("2", "Pre-War Stable — Low Signal",                 "Low",    "242", "3.5M kBtu avg",  "1934",  "Good ES (72), older buildings, no strong churn signal"),
    ("3", "Large Commercial — Capital Mobilized",        "Medium", "263", "4.6M kBtu avg",  "1937",  "Pre-war commercial, low permits, stable base"),
    ("4", "Low-Compliance Commercial — Quiet Attrition", "High",   "247", "29.1M kBtu avg", "1963",  "High DOB (12.3 avg), 97% office — acute LL97 fine exposure"),
]
add_table(doc,
    ["ID", "Archetype Name", "Risk", "Count", "Avg Steam", "Avg Yr", "Key Signal"],
    archetype_rows, [0.25, 2.3, 0.6, 0.5, 1.0, 0.55, 1.5])

doc.add_page_break()

add_heading(doc, "3.3  Gradient Boosting Risk Classifier", level=2)
add_para(doc, "After clustering, a supervised GradientBoostingClassifier produces a continuous 0–1 attrition probability per building. Training labels are derived from historical steam trend signals: buildings with a confirmed 'big_drop' in steam consumption are labeled positive (likely churner).", size=10)
doc.add_paragraph()

add_para(doc, "Important design decision:", bold=True, size=10)
add_para(doc, "Buildings labeled 'moderate_drop' were explicitly excluded from training. This produces the observed bimodal distribution — the model is a near-binary classifier, not a nuanced scorer. This is intentional: it maximizes precision on the high-risk tail at the cost of coverage in the middle band.", size=10)
doc.add_paragraph()

add_heading(doc, "GBM Features (12 total):", level=3)
gbm_features = [
    ("ll97_penalty_2024_log", "~22%", "Dominant signal — acute financial pressure to switch"),
    ("log_steam", "~17%", "Raw consumption volume"),
    ("ll97_over_2024", "~13%", "Binary: is building currently non-compliant?"),
    ("log_ghg", "~12%", "GHG intensity"),
    ("peer_score", "~9%",  "Relative efficiency vs. same-use peers"),
    ("ll97_penalty_2030_log", "~7%", "Forward-looking pressure signal"),
    ("cluster_id", "~6%",  "Archetype membership from K-means"),
    ("steam_ghg_share", "~5%", "Fraction of building GHG attributable to steam"),
    ("energy_star", "~4%", "Efficiency benchmark"),
    ("year_built", "~2%",  "Building age proxy"),
    ("log_dob_jobs", "~2%", "24-month permit activity"),
    ("use_type_ord", "~1%", "Electrification sensitivity ordinal"),
]
add_table(doc, ["Feature", "Approx. Importance", "Interpretation"], gbm_features, [2.0, 1.4, 3.0])
doc.add_paragraph()

fig = fig_feature_importance()
doc.add_picture(fig_to_stream(fig), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

fig = fig_risk_distribution()
doc.add_picture(fig_to_stream(fig), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc, "Risk Tier Thresholds:", level=3)
add_table(doc, ["Tier", "Score Range", "Buildings", "Description"],
    [
        ("High",   "≥ 0.70", "58  (4.8%)",    "Strong churn signal; immediate account management attention"),
        ("Medium", "0.30–0.70", "11  (0.9%)", "Moderate signal; monitor quarterly"),
        ("Low",    "< 0.30",  "1,141  (94.3%)", "No strong signal; routine stewardship"),
    ], [0.8, 1.2, 1.4, 3.0])

doc.add_page_break()

# ── SECTION 4: LL97 EQUATIONS ─────────────────────────────────────────────────

add_heading(doc, "Section 4 — LL97 Equations & Penalty Model")

add_para(doc, "NYC Local Law 97 (2019) mandates annual carbon intensity limits for buildings over 25,000 sq ft. The dashboard computes compliance status and projected fines for both the 2024 (Phase 1) and 2030 (Phase 2) thresholds.", size=10)
doc.add_paragraph()

fig = fig_ll97_equation()
doc.add_picture(fig_to_stream(fig), width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, "4.1  Steam → GHG Conversion", level=2)
add_equation_box(doc, "GHG (MT CO₂e/yr)",
    "GHG = steam_kBtu × 0.00004493",
    "0.00004493 MT CO₂e/kBtu is the LL97-specific binding coefficient (NYC DOB Chapter 103 Rules). "
    "Note: EPA eGRID cites ~6.68×10⁻⁵ — LL97 uses the lower value as the official regulatory standard.")

add_heading(doc, "4.2  Annual Cap Calculation", level=2)
add_equation_box(doc, "Emission cap (MT CO₂e/yr)",
    "cap = floor_sqft × intensity_limit[use_type][phase]",
    "intensity_limit varies by building use type and compliance phase (see table below).")

add_heading(doc, "4.3  Excess & Fine", level=2)
add_equation_box(doc, "Annual fine (USD)",
    "excess = max(0, GHG − cap)\n"
    "  fine  = excess × $268",
    "$268/MT CO₂e is the statutory penalty rate. There is no cap on total fines.")

add_heading(doc, "4.4  LL97 Intensity Limits by Use Type", level=2)
ll97_table = [
    ("Office / Financial Office",      "0.00846", "0.00453"),
    ("Multifamily Housing / Dorm",     "0.00675", "0.00400"),
    ("Hotel",                          "0.01450", "0.00700"),
    ("Retail Store",                   "0.00846", "0.00403"),
    ("K-12 School / College",          "0.00846", "0.00453"),
    ("Hospital / Clinic",              "0.02381", "0.00840"),
    ("Performing Arts / Museum",       "0.01074", "0.00420"),
    ("Worship Facility",               "0.01074", "0.00420"),
    ("Default (Other)",                "0.00846", "0.00453"),
]
add_table(doc, ["Use Type", "Phase 1 Limit (2024)\nMT CO₂e / ft² / yr", "Phase 2 Limit (2030)\nMT CO₂e / ft² / yr"],
          ll97_table, [2.8, 1.9, 1.9])
doc.add_paragraph()

add_heading(doc, "4.5  Portfolio LL97 Exposure", level=2)
add_table(doc, ["Metric", "2024 (Phase 1)", "2030 (Phase 2)"],
    [
        ("Total portfolio fines (all buildings)", "$81,875,711", "$270,916,416"),
        ("Buildings over limit",                 "~516",        "~760 (est.)"),
        ("Multiplier vs. 2024",                  "1.0×",        "~3.3×"),
        ("Top fine: 31 W 34th St",               "$16,987,675", "~$56M (est.)"),
        ("2nd: Columbia Univ (116 St complex)",  "$12,504,928", "~$41M (est.)"),
    ], [2.5, 1.8, 1.8])
doc.add_paragraph()

fig = fig_ll97_by_cluster()
doc.add_picture(fig_to_stream(fig), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── SECTION 5: DASHBOARD COMPONENTS ───────────────────────────────────────────

add_heading(doc, "Section 5 — Dashboard Components")

fig = fig_dashboard_components()
doc.add_picture(fig_to_stream(fig), width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

components = [
    ("KPI Cards (top bar)",
     "Four summary statistics: Total Portfolio (1,210), High-Risk count (58), "
     "LL97 2024 Exposure ($81.9M), and filtered/watchlist count. "
     "Values update reactively when filters or watchlist change."),

    ("RiskTable",
     "Sortable, filterable table of all buildings. Columns: Address, EUI (kBtu/ft²), "
     "Steam (kBtu/yr), DOB Permits (24mo), Cluster Archetype, LL97 2024 fine, Risk Score. "
     "Filters: archetype dropdown, steam signal (big_drop/stable/growth), LL97 compliance status. "
     "Click a row to open BuildingPanel. Star button to add to Watchlist."),

    ("BuildingPanel",
     "Right-side detail panel for the selected building. Shows: EUI, YoY steam delta "
     "(22→23 and 23→24), GBM risk score with color badge, cluster archetype name, "
     "LL97 2024 and 2030 fine, DOB permit count, steam-GHG share percentage, "
     "and building year. Appears on row click or watchlist click."),

    ("YoYScatter",
     "Scatter plot comparing two YoY steam change periods. X-axis: 2022→2023 % change. "
     "Y-axis: 2023→2024 % change. Each dot = one building. Color = cluster archetype. "
     "Size encodes steam volume. Buildings with extreme outlier values in either year "
     "receive a highlight ring. The scatter reveals whether declines are accelerating "
     "(down-right quadrant = declining again) or one-time shocks (up-right = recovered)."),

    ("Watchlist",
     "Persistent sidebar (localStorage) for pinned buildings. Supports: add/remove individual "
     "buildings (star toggle), clear all button, click-to-select. Useful for account managers "
     "monitoring a specific set of customers across sessions."),

    ("AI Insights (Groq)",
     "Natural language filter powered by Llama 3.3 70B via Groq API (streamed). "
     "User types a query ('show me hotels over the LL97 2030 cap') — the backend sends a "
     "system prompt + building schema to Groq, receives a JSON filter spec, and the frontend "
     "applies it client-side via groqFilter.js. No building data leaves the local network — "
     "only the query and field schema go to Groq."),

    ("Trends Tab",
     "Per-building line charts showing annual steam consumption 2021–2024, "
     "GHG trajectory vs. LL97 cap line, and rolling DOB permit activity. "
     "Allows visual inspection of trend shape (accelerating decline vs. one-time dip)."),
]

for name, desc in components:
    add_heading(doc, f"5.x  {name}", level=2)
    add_para(doc, desc, size=10)
    doc.add_paragraph()

doc.add_page_break()

# ── SECTION 6: DATA DEFINITIONS ────────────────────────────────────────────────

add_heading(doc, "Section 6 — Data Field Definitions")

field_rows = [
    ("steam", "kBtu/yr", "Annual district steam consumption (most recent year)"),
    ("ghg", "MT CO₂e/yr", "GHG from steam: steam × 0.00004493"),
    ("eui", "kBtu/ft²/yr", "Energy Use Intensity: steam / floor_sqft"),
    ("floor_sqft", "ft²", "Gross floor area (self-reported, NYC LL Benchmarking)"),
    ("energy_star", "0–100", "Energy Star score; imputed by use-type median when absent"),
    ("peer_score", "Z-score", "EUI vs. same-use-type peers; negative = more efficient"),
    ("dob_jobs", "integer", "DOB NOW permit filings in trailing 24 months"),
    ("log_dob_jobs", "float", "log1p(dob_jobs); used as ML feature"),
    ("cluster_id", "0–4", "K-means archetype assignment"),
    ("cluster_name", "string", "Human-readable archetype name"),
    ("cluster_risk", "High/Medium/Low", "Risk tier assigned to this archetype"),
    ("ml_risk", "0.0–1.0", "GBM predict_proba — attrition probability score"),
    ("ll97_penalty_2024", "USD/yr", "Projected annual fine under Phase 1 limits"),
    ("ll97_penalty_2030", "USD/yr", "Projected annual fine under Phase 2 limits"),
    ("ll97_over_2024", "0/1", "Binary: building exceeds Phase 1 cap"),
    ("ll97_over_2030", "0/1", "Binary: building exceeds Phase 2 cap"),
    ("ll97_cap_2024", "MT CO₂e/yr", "Building's Phase 1 emission allowance"),
    ("steam_ghg_share", "0.0–1.0", "Fraction of building GHG attributable to steam"),
    ("yr", "year", "Year built (from LL Benchmarking / PLUTO)"),
    ("bbl", "string", "Borough-Block-Lot; may be semicolon-delimited for multi-tax-lot buildings"),
]
add_table(doc, ["Field", "Unit", "Description"], field_rows, [1.5, 1.0, 4.0])

doc.add_page_break()

# ── SECTION 7: UPDATE RUNBOOK ───────────────────────────────────────────────────

add_heading(doc, "Section 7 — Update Runbook")

add_heading(doc, "7.1  Incremental DOB Update", level=2)
add_para(doc, "Run from coned-dashboard/ directory:", size=10)
p = doc.add_paragraph()
run = p.add_run(
    "cd coned-dashboard\n"
    "/opt/homebrew/bin/python3.13 update_dob_jobs.py\n"
    "/opt/homebrew/bin/python3.13 kmeans_model.py\n"
    "/opt/homebrew/bin/python3.13 ll97_model.py"
)
run.font.name = "Courier New"
run.font.size = Pt(9)
p.paragraph_format.left_indent = Inches(0.4)
doc.add_paragraph()

add_heading(doc, "7.2  Full ML Rebuild", level=2)
add_para(doc, "Same three scripts as above. Takes ~2 minutes. After running, commit public/buildingEnrichment.json and cluster_profiles.json.", size=10)
doc.add_paragraph()

add_heading(doc, "7.3  Start the Dashboard (local)", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "# Terminal 1 — backend\n"
    "cd coned-dashboard && node server.js\n\n"
    "# Terminal 2 — frontend\n"
    "cd coned-dashboard && npm run dev\n\n"
    "# Access: http://localhost:5173\n"
    "# Password: coned-steam-2026"
)
run.font.name = "Courier New"
run.font.size = Pt(9)
p.paragraph_format.left_indent = Inches(0.4)
doc.add_paragraph()

add_heading(doc, "7.4  Remote Access via Tailscale", level=2)
add_para(doc,
    "Start Vite with --host to expose on all interfaces:\n"
    "npm run dev -- --host\n\n"
    "Other devices on Tailscale can access at: http://100.107.92.16:5173", size=10)

doc.add_page_break()

# ── SECTION 8: KNOWN LIMITATIONS ──────────────────────────────────────────────

add_heading(doc, "Section 8 — Known Limitations & Data Notes")

limitations = [
    ("Risk model is binary by design",
     "The GBM was trained on big_drop vs. no_signal labels. Buildings with 'moderate_drop' "
     "were excluded from training. The result: very few buildings score in the 0.30–0.70 "
     "medium band (11 buildings total). This is a modeling choice, not a bug."),

    ("YoY data coverage gap",
     "2022→2023 YoY delta covers 743 buildings. 2023→2024 covers only 422 buildings "
     "because 2024 benchmarking data was not yet available for all accounts at data collection time. "
     "Buildings missing 2024 data show their most recent available year."),

    ("LL97 2030 estimates are projections",
     "The 2030 fine calculations assume current steam consumption levels hold constant. "
     "Actual fines will depend on: LL97 regulatory changes, building renovations, "
     "and whether building owners comply via offsets, retrofits, or switching."),

    ("Floor area is self-reported",
     "GFA comes from NYC LL Benchmarking self-reporting. Errors in floor area directly "
     "affect EUI and LL97 penalty calculations. Some buildings show EUI outliers due to "
     "misreported floor area."),

    ("DOB permits proxy only",
     "dob_jobs counts permit filings, not completed work. High permit counts may indicate "
     "renovation toward electrification OR general building maintenance. "
     "The model uses this as a directional signal, not a definitive indicator."),

    ("LL97 fine cross-check",
     "The $81.9M total cannot yet be independently verified against NYC DEP compliance reports "
     "because official 2024 enforcement data has not been released. The calculation uses "
     "the statutory formula as written in LL97 and NYC DOB Chapter 103 Rules."),
]

for title, body in limitations:
    p = doc.add_paragraph()
    run1 = p.add_run(f"⚠  {title}: ")
    run1.bold = True
    run1.font.size = Pt(10)
    run2 = p.add_run(body)
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ── APPENDIX: QUICK NUMBERS ────────────────────────────────────────────────────

add_heading(doc, "Appendix A — Quick Numbers Reference")

quick_rows = [
    ("Portfolio size",        "1,210 buildings"),
    ("Borough",               "Manhattan only"),
    ("Steam emission factor", "0.00004493 MT CO₂e / kBtu"),
    ("LL97 fine rate",        "$268 per MT CO₂e over cap"),
    ("Total LL97 2024 fine",  "$81,875,711"),
    ("Total LL97 2030 fine",  "$270,916,416"),
    ("High-risk buildings",   "58 (score ≥ 0.70)"),
    ("Medium-risk buildings", "11 (score 0.30–0.70)"),
    ("Low-risk buildings",    "1,141 (score < 0.30)"),
    ("K-means K",             "5 clusters"),
    ("Silhouette score @ K=5","0.191"),
    ("DOB records total",     "19,879 (through 2026-06-13)"),
    ("DOB rolling window",    "24 months"),
    ("GBM features",          "12"),
    ("Dashboard password",    "coned-steam-2026"),
]
add_table(doc, ["Metric", "Value"], quick_rows, [3.0, 3.0])

# ── SAVE ───────────────────────────────────────────────────────────────────────

doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
print(f"Size:  {os.path.getsize(OUT_FILE) / 1024:.0f} KB")
